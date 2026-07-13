from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from services.helpers import set_font


def normalize_line(value):
    """
    Menghapus spasi berlebih dalam satu baris.
    """
    if value is None:
        return ""

    return " ".join(str(value).split())


def normalize_multiline_text(value):
    """
    Membersihkan teks beberapa baris.

    Fungsi ini:
    - menyeragamkan jenis Enter;
    - menghapus Enter kosong;
    - menghapus spasi berlebih;
    - mempertahankan baris alamat yang benar.
    """
    if value is None:
        return ""

    text = str(value)
    text = text.replace("\r\n", "\n")
    text = text.replace("\r", "\n")

    clean_lines = []

    for line in text.split("\n"):
        clean_line = normalize_line(line)

        if clean_line:
            clean_lines.append(clean_line)

    return "\n".join(clean_lines)


def set_cell_margins(
    cell,
    top=50,
    start=70,
    bottom=50,
    end=70,
):
    """
    Mengatur margin internal sel tabel.

    Nilai menggunakan satuan twip/dxa Word.
    """
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    tc_mar = tc_pr.first_child_found_in("w:tcMar")

    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    margins = {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }

    for margin_name, margin_value in margins.items():
        tag = qn(f"w:{margin_name}")
        node = tc_mar.find(tag)

        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)

        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(
    cell,
    text,
    size=8,
    bold=False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
):
    """
    Memasukkan teks ke sel tabel tanpa menghasilkan jarak kosong
    antarbaris alamat.
    """
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    paragraph = cell.paragraphs[0]
    paragraph.alignment = align

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1

    clean_text = normalize_multiline_text(text)

    if clean_text:
        lines = clean_text.split("\n")
    else:
        lines = [""]

    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_font(run, size=size, bold=bold)

        # Membuat pindah baris biasa, bukan paragraf baru.
        if index < len(lines) - 1:
            run.add_break()

    set_cell_margins(cell)


def set_col_widths(table):
    """
    Mengatur lebar setiap kolom tabel.
    """
    widths = [
        Cm(1.0),   # No Urut
        Cm(3.2),   # No Surat
        Cm(10.0),  # Ditujukan
        Cm(2.2),   # No X5
        Cm(1.8),   # Berat
        Cm(2.0),   # Tarif
        Cm(3.0),   # Keterangan
    ]

    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = width


def set_signature_paragraph(
    paragraph,
    text,
    size=9,
    bold=False,
):
    """
    Mengatur paragraf tanda tangan agar tidak memiliki jarak tambahan.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1

    run = paragraph.add_run(str(text or ""))
    set_font(run, size=size, bold=bold)


def add_signature_blank_line(cell):
    """
    Menambahkan baris kosong untuk area tanda tangan.
    """
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1

    run = paragraph.add_run("")
    set_font(run, size=9)


def build_daftar_pengiriman(data):
    """
    Membuat dokumen Daftar Pengiriman Pos Surat Dinas (X5).
    """
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = (
        section.page_height,
        section.page_width,
    )

    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # =====================================================
    # JUDUL
    # =====================================================

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    title_run = title.add_run(
        "DAFTAR PENGIRIMAN POS SURAT DINAS (X5)"
    )
    set_font(title_run, size=11, bold=True)

    # =====================================================
    # TABEL
    # =====================================================

    rows_data = data.get("selected_surat", [])

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.autofit = False

    headers = [
        "No\nUrut",
        "No. Surat",
        "Ditujukan",
        "No.X5",
        "Berat (Gr)",
        "Tarif (Rp.)",
        "Keterangan",
    ]

    for index, header in enumerate(headers):
        set_cell_text(
            cell=table.rows[0].cells[index],
            text=header,
            size=8,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for row_number, row_data in enumerate(rows_data, start=1):
        cells = table.add_row().cells

        nomor_surat = normalize_line(
            row_data.get("nomor_referensi", "")
        )

        alamat = normalize_multiline_text(
            row_data.get("alamat", "")
        )

        set_cell_text(
            cell=cells[0],
            text=f"{row_number}.",
            size=8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

        set_cell_text(
            cell=cells[1],
            text=nomor_surat,
            size=8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

        set_cell_text(
            cell=cells[2],
            text=alamat,
            size=7,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )

        set_cell_text(
            cell=cells[3],
            text="",
            size=8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

        set_cell_text(
            cell=cells[4],
            text="",
            size=8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

        set_cell_text(
            cell=cells[5],
            text="",
            size=8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

        set_cell_text(
            cell=cells[6],
            text="",
            size=8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    set_col_widths(table)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(8)

    # =====================================================
    # TANDA TANGAN
    # =====================================================

    signature_table = doc.add_table(rows=1, cols=2)
    signature_table.autofit = False

    left_cell = signature_table.cell(0, 0)
    right_cell = signature_table.cell(0, 1)

    left_cell.width = Cm(12)
    right_cell.width = Cm(12)

    left_cell.text = ""
    right_cell.text = ""

    # Kolom kiri: pengirim
    set_signature_paragraph(
        paragraph=left_cell.paragraphs[0],
        text="Pengirim,",
        size=9,
    )

    for _ in range(3):
        add_signature_blank_line(left_cell)

    name_paragraph = left_cell.add_paragraph()
    set_signature_paragraph(
        paragraph=name_paragraph,
        text=data.get("petugas_nama", ""),
        size=9,
        bold=True,
    )

    nip_paragraph = left_cell.add_paragraph()
    set_signature_paragraph(
        paragraph=nip_paragraph,
        text=data.get("petugas_nip", ""),
        size=8,
    )

    # Kolom kanan: petugas pos
    set_signature_paragraph(
        paragraph=right_cell.paragraphs[0],
        text=f"Jakarta, {data.get('tanggal_surat', '')}",
        size=9,
    )

    postal_title = right_cell.add_paragraph()
    set_signature_paragraph(
        paragraph=postal_title,
        text="Petugas Pos",
        size=9,
    )

    for _ in range(2):
        add_signature_blank_line(right_cell)

    postal_name = right_cell.add_paragraph()
    set_signature_paragraph(
        paragraph=postal_name,
        text=data.get("petugas_pos_nama", ""),
        size=9,
        bold=True,
    )

    postal_nip = right_cell.add_paragraph()
    set_signature_paragraph(
        paragraph=postal_nip,
        text=data.get("petugas_pos_nip", ""),
        size=8,
    )

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output