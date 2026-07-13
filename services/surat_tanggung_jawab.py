from io import BytesIO

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from services.helpers import angka_ke_kata


def set_font(run, size=12, bold=False, underline=False, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.0):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_spacer(doc, height_pt=12):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=height_pt, line_spacing=1.0)
    r = p.add_run("")
    set_font(r, size=1)
    return p


def remove_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)

        border.set(qn("w:val"), "nil")


def set_table_indent(table, indent_cm):
    tbl_pr = table._tbl.tblPr

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)

    tbl_ind.set(qn("w:w"), str(int(indent_cm * 567)))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()

    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)

    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, size=12, bold=False, underline=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)

    r = p.add_run(str(text))
    set_font(r, size=size, bold=bold, underline=underline)


def add_identity_table(doc, data):
    nama = data.get("nama_penanggung_jawab", "")
    email = data.get("email_penanggung_jawab", "")
    jabatan = data.get("jabatan_penanggung_jawab", "")
    jumlah = data.get("jumlah_dokumen", "0")

    nama_display = f"{nama} ({email})" if email else nama
    jumlah_display = f"{jumlah} ({angka_ke_kata(jumlah)}) buah"

    rows = [
        ("Nama", ":", nama_display),
        ("Jabatan", ":", jabatan),
        ("Jumlah", ":", jumlah_display),
    ]

    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    remove_table_borders(table)
    set_table_indent(table, 0.95)

    widths = [2.4, 0.4, 11.2]

    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            set_cell_width(cell, widths[col_idx])
            set_cell_text(cell, value, size=12, align=WD_ALIGN_PARAGRAPH.LEFT)

    return table


def add_signature_cell(cell, title, name, nip_text):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ""

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=54, line_spacing=1.0)
    r = p.add_run(title)
    set_font(r, size=12)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
    r = p.add_run(name)
    set_font(r, size=12, underline=True)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
    r = p.add_run(nip_text)
    set_font(r, size=12)


def build_surat_tanggung_jawab(data):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    # Mengikuti file contoh: margin kiri/kanan sekitar 1 inch,
    # atas sekitar 1 inch, bawah sedikit lebih kecil.
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(1.75)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    # Judul
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=30, line_spacing=1.0)
    r = p.add_run("SURAT KETERANGAN TANGGUNG JAWAB")
    set_font(r, size=12, bold=True)

    # Pembuka
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, before=0, after=8, line_spacing=1.0)
    r = p.add_run("Yang bertanda tangan dibawah ini :")
    set_font(r, size=12)

    # Identitas
    add_identity_table(doc, data)

    add_spacer(doc, height_pt=10)

    # Isi paragraf pertama
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, before=0, after=8, line_spacing=1.5)
    r = p.add_run(
        "Dengan ini menyatakan bertanggungjawab atas penggunaan Carik Surat Terdaftar (X-5) "
        "yang digunakan untuk keperluan Surat Dinas dan Berkas Perkara pada Kepaniteraan "
        "Mahkamah Agung RI."
    )
    set_font(r, size=12)

    # Isi paragraf kedua
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, before=0, after=28, line_spacing=1.5)
    r = p.add_run(
        "Penerimaan carik dengan buku tanda terima dan sebulan akhir Dinas Kantor Pos "
        "Mahkamah Agung RI carik tersebut sudah diterima kembali sesuai dengan jumlah "
        "yang diterima."
    )
    set_font(r, size=12)

    # Tanggal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, before=0, after=20, line_spacing=1.0)
    r = p.add_run(f"{data.get('kota', 'Jakarta')}, {data.get('tanggal_surat', '')}")
    set_font(r, size=12)

    # Tanda tangan dua kolom
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    set_cell_width(left_cell, 8.0)
    set_cell_width(right_cell, 8.0)

    add_signature_cell(
        left_cell,
        "Petugas Pos,",
        data.get("petugas_pos_nama", ""),
        data.get("petugas_pos_nip", "")
    )

    add_signature_cell(
        right_cell,
        "Petugas",
        data.get("petugas_nama", ""),
        data.get("petugas_nip", "")
    )

    # Jarak menuju bagian Mengetahui
    add_spacer(doc, height_pt=95)

    # Mengetahui
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
    r = p.add_run("Mengetahui :")
    set_font(r, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
    r = p.add_run(data.get("mengetahui_jabatan_1", ""))
    set_font(r, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=58, line_spacing=1.0)
    r = p.add_run(data.get("mengetahui_jabatan_2", ""))
    set_font(r, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
    r = p.add_run(data.get("mengetahui_nama", ""))
    set_font(r, size=12, underline=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
    r = p.add_run(data.get("mengetahui_nip", ""))
    set_font(r, size=12)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    return bio